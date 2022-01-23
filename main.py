try:
    import os
    import re
    import sys
    import tempfile
    from collections import Counter

    import fitz
    from PyQt5.QtCore import Qt, QRectF, pyqtSignal, QT_VERSION_STR, QObject
    from PyQt5.QtCore import pyqtSlot
    from PyQt5.QtGui import QImage, QPixmap, QPainterPath
    from PyQt5.QtGui import QIntValidator
    from PyQt5.QtGui import QStandardItemModel
    from PyQt5.QtWidgets import QApplication, QWidget, QPushButton, \
        QGroupBox, QFormLayout, QVBoxLayout, QComboBox, QLabel, \
        QHBoxLayout, QLineEdit, QGridLayout, QTreeView, \
        QSplitter, QSizePolicy, QFileDialog, QAbstractItemView, \
        QTextBrowser, QDialog, QProgressBar, QLayout
    from PyQt5.QtWidgets import QGraphicsView, QGraphicsScene, QTableView, QHeaderView
    from PyQt5.QtCore import QThread
    from enum import Enum
    import time

    from pdfminer.converter import PDFPageAggregator
    from pdfminer.layout import *
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
    from pdfminer.pdfpage import PDFPage
    from pdfminer.pdfpage import PDFTextExtractionNotAllowed
    from pdfminer.pdfparser import PDFParser
    import traceback
    from docx import Document

    # hide pymupdf errors output
    fitz.TOOLS.mupdf_display_errors(False)
    fitz.TOOLS.mupdf_display_warnings(False)
except:
    import traceback

    print(traceback.format_exc())
    sys.exit(-1)


# refer: https://github.com/marcel-goldschen-ohm/PyQtImageViewer/blob/master/QtImageViewer.py
class QtImageViewer(QGraphicsView):
    """ PyQt image viewer widget for a QPixmap in a QGraphicsView scene with mouse zooming and panning.

    Displays a QImage or QPixmap (QImage is internally converted to a QPixmap).
    To display any other image format, you must first convert it to a QImage or QPixmap.

    Some useful image format conversion utilities:
        qimage2ndarray: NumPy ndarray <==> QImage    (https://github.com/hmeine/qimage2ndarray)
        ImageQt: PIL Image <==> QImage  (https://github.com/python-pillow/Pillow/blob/master/PIL/ImageQt.py)

    Mouse interaction:
        Left mouse button drag: Pan image.
        Right mouse button drag: Zoom box.
        Right mouse button doubleclick: Zoom to show entire image.
    """

    # Mouse button signals emit image scene (x, y) coordinates.
    # !!! For image (row, column) matrix indexing, row = y and column = x.
    leftMouseButtonPressed = pyqtSignal(float, float)
    rightMouseButtonPressed = pyqtSignal(float, float)
    leftMouseButtonReleased = pyqtSignal(float, float)
    rightMouseButtonReleased = pyqtSignal(float, float)
    leftMouseButtonDoubleClicked = pyqtSignal(float, float)
    rightMouseButtonDoubleClicked = pyqtSignal(float, float)

    def __init__(self):
        QGraphicsView.__init__(self)

        # Image is displayed as a QPixmap in a QGraphicsScene attached to this QGraphicsView.
        self.scene = QGraphicsScene()
        self.setScene(self.scene)

        # Store a local handle to the scene's current image pixmap.
        self._pixmapHandle = None

        # Image aspect ratio mode.
        # !!! ONLY applies to full image. Aspect ratio is always ignored when zooming.
        #   Qt.IgnoreAspectRatio: Scale image to fit viewport.
        #   Qt.KeepAspectRatio: Scale image to fit inside viewport, preserving aspect ratio.
        #   Qt.KeepAspectRatioByExpanding: Scale image to fill the viewport, preserving aspect ratio.
        self.aspectRatioMode = Qt.KeepAspectRatio

        # Scroll bar behaviour.
        #   Qt.ScrollBarAlwaysOff: Never shows a scroll bar.
        #   Qt.ScrollBarAlwaysOn: Always shows a scroll bar.
        #   Qt.ScrollBarAsNeeded: Shows a scroll bar only when zoomed.
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)

        # Stack of QRectF zoom boxes in scene coordinates.
        self.zoomStack = []

        # Flags for enabling/disabling mouse interaction.
        self.canZoom = True
        self.canPan = True

    def hasImage(self):
        """ Returns whether or not the scene contains an image pixmap.
        """
        return self._pixmapHandle is not None

    def clearImage(self):
        """ Removes the current image pixmap from the scene if it exists.
        """
        if self.hasImage():
            self.scene.removeItem(self._pixmapHandle)
            self._pixmapHandle = None

    def pixmap(self):
        """ Returns the scene's current image pixmap as a QPixmap, or else None if no image exists.
        :rtype: QPixmap | None
        """
        if self.hasImage():
            return self._pixmapHandle.pixmap()
        return None

    def image(self):
        """ Returns the scene's current image pixmap as a QImage, or else None if no image exists.
        :rtype: QImage | None
        """
        if self.hasImage():
            return self._pixmapHandle.pixmap().toImage()
        return None

    def setImage(self, image):
        """ Set the scene's current image pixmap to the input QImage or QPixmap.
        Raises a RuntimeError if the input image has type other than QImage or QPixmap.
        :type image: QImage | QPixmap
        """
        if type(image) is QPixmap:
            pixmap = image
        elif type(image) is QImage:
            pixmap = QPixmap.fromImage(image)
        else:
            raise RuntimeError("ImageViewer.setImage: Argument must be a QImage or QPixmap.")
        if self.hasImage():
            self._pixmapHandle.setPixmap(pixmap)
        else:
            self._pixmapHandle = self.scene.addPixmap(pixmap)
        self.setSceneRect(QRectF(pixmap.rect()))  # Set scene size to image size.
        self.updateViewer()

    def loadImageFromFile(self, fileName=""):
        """ Load an image from file.
        Without any arguments, loadImageFromFile() will popup a file dialog to choose the image file.
        With a fileName argument, loadImageFromFile(fileName) will attempt to load the specified image file directly.
        """
        if len(fileName) == 0:
            if QT_VERSION_STR[0] == '4':
                fileName = QFileDialog.getOpenFileName(self, "Open image file.")
            elif QT_VERSION_STR[0] == '5':
                fileName, dummy = QFileDialog.getOpenFileName(self, "Open image file.")
        if len(fileName) and os.path.isfile(fileName):
            image = QImage(fileName)
            self.setImage(image)

    def updateViewer(self):
        """ Show current zoom (if showing entire image, apply current aspect ratio mode).
        """
        if not self.hasImage():
            return
        if len(self.zoomStack) and self.sceneRect().contains(self.zoomStack[-1]):
            self.fitInView(self.zoomStack[-1], Qt.IgnoreAspectRatio)  # Show zoomed rect (ignore aspect ratio).
        else:
            self.zoomStack = []  # Clear the zoom stack (in case we got here because of an invalid zoom).
            self.fitInView(self.sceneRect(), self.aspectRatioMode)  # Show entire image (use current aspect ratio mode).

    def resizeEvent(self, event):
        """ Maintain current zoom on resize.
        """
        self.updateViewer()

    def mousePressEvent(self, event):
        """ Start mouse pan or zoom mode.
        """
        scenePos = self.mapToScene(event.pos())
        if event.button() == Qt.LeftButton:
            if self.canPan:
                self.setDragMode(QGraphicsView.ScrollHandDrag)
            self.leftMouseButtonPressed.emit(scenePos.x(), scenePos.y())
        elif event.button() == Qt.RightButton:
            if self.canZoom:
                self.setDragMode(QGraphicsView.RubberBandDrag)
            self.rightMouseButtonPressed.emit(scenePos.x(), scenePos.y())
        QGraphicsView.mousePressEvent(self, event)

    def mouseReleaseEvent(self, event):
        """ Stop mouse pan or zoom mode (apply zoom if valid).
        """
        QGraphicsView.mouseReleaseEvent(self, event)
        scenePos = self.mapToScene(event.pos())
        if event.button() == Qt.LeftButton:
            self.setDragMode(QGraphicsView.NoDrag)
            self.leftMouseButtonReleased.emit(scenePos.x(), scenePos.y())
        elif event.button() == Qt.RightButton:
            if self.canZoom:
                viewBBox = self.zoomStack[-1] if len(self.zoomStack) else self.sceneRect()
                selectionBBox = self.scene.selectionArea().boundingRect().intersected(viewBBox)
                self.scene.setSelectionArea(QPainterPath())  # Clear current selection area.
                if selectionBBox.isValid() and (selectionBBox != viewBBox):
                    self.zoomStack.append(selectionBBox)
                    self.updateViewer()
            self.setDragMode(QGraphicsView.NoDrag)
            self.rightMouseButtonReleased.emit(scenePos.x(), scenePos.y())

    def mouseDoubleClickEvent(self, event):
        """ Show entire image.
        """
        scenePos = self.mapToScene(event.pos())
        if event.button() == Qt.LeftButton:
            self.leftMouseButtonDoubleClicked.emit(scenePos.x(), scenePos.y())
        elif event.button() == Qt.RightButton:
            if self.canZoom:
                self.zoomStack = []  # Clear zoom stack.
                self.updateViewer()
            self.rightMouseButtonDoubleClicked.emit(scenePos.x(), scenePos.y())
        QGraphicsView.mouseDoubleClickEvent(self, event)


class FileInfo(object):
    suffix = ''

    def __init__(self, path=None):
        self.path = path

    def pages(self):
        return 'Unknown'

    def basename(self):
        if self.path:
            return os.path.basename(self.path)

    def fullpath(self):
        if self.path:
            return os.path.abspath(self.path)

    def filetype(self):
        if self.path:
            _, ext = os.path.splitext(self.path)
            return ext.lower()


class PDFInfo(FileInfo):
    suffix = '.pdf'

    def pages(self):
        if self.path and os.path.exists(self.path):
            doc = fitz.open(self.path)
            return doc.page_count

    @staticmethod
    def is_pdf(filename):
        fi = FileInfo(filename)
        ext = fi.filetype()
        return ext == PDFInfo.suffix


class WatermarkGuesser(object):

    def __init__(self, doc=None):
        self.doc = doc
        self.most_common = 10

    def guess(self):
        if self.doc is None:
            return
        all_data = []
        for page in self.doc.pages():
            text = page.get_text()
            data = text.split('\n')
            data = list(filter(lambda x: x.strip(), data))
            all_data.extend(data)
        c = Counter(all_data)
        items = c.most_common(self.most_common)
        return items

    def __call__(self, *args, **kwargs):
        return self.guess()


class ImageGuesser(WatermarkGuesser):

    def __init__(self, doc=None):
        super().__init__(doc)
        self.name_ref_map = {}

    class Image(object):
        def __init__(self, ref, width=0, height=0):
            self.ref = ref
            self.width = width
            self.height = height

        def __str__(self):
            return f'Image {self.ref}'

        def __hash__(self):
            return hash(str(self))

    def guess0(self):
        if self.doc is None:
            return

        all_data = []
        for page in self.doc.pages():
            image_list = page.getImageList()
            ref = list(map(lambda x: x[0], image_list))
            # get width and height
            for item in ref:
                d = self.doc.extract_image(item)

                im = ImageGuesser.Image(item, d['width'], d['height'])
                all_data.append(str(im))
                self.name_ref_map[str(im)] = str(item)

        c = Counter(all_data)
        items = c.most_common(self.most_common)
        return items

    def guess(self):
        if self.doc is None:
            return

        all_data = []
        for page in self.doc.pages():
            page.cleanContents()  # cleanup page painting commands
            xref = page.getContents()[0]  # get xref of the resulting source
            cont0 = self.doc.xrefStream(xref).decode()  # .splitlines()
            image_ids = re.findall(r'/Ima?g?e?(\d+) Do', cont0)
            for ele in image_ids:
                im = ImageGuesser.Image(ele)
                self.name_ref_map[str(im)] = ele
                all_data.append(str(im))

        c = Counter(all_data)
        items = c.most_common(self.most_common)
        return items


class GuesserType(Enum):
    Text = 1
    Image = 2


class GuesserThread(QThread):
    most_common = 10

    sinProgress = pyqtSignal(str, GuesserType, int, int)  # filename, GuesserType, current, total
    # sinDone = pyqtSignal(str)  # filename
    sinResult = pyqtSignal(str, list)  # filename, [str]

    def __init__(self, parent=None):
        super(GuesserThread, self).__init__(parent)
        self._filename = None

    @property
    def filename(self):
        return self._filename

    @filename.setter
    def filename(self, v):
        self._filename = v

    def run(self):
        raise NotImplementedError


class TextGuesserThread(GuesserThread):

    def run(self):
        doc = fitz.open(self.filename)
        all_data = []
        pages_count = doc.page_count
        for idx, page in enumerate(doc.pages(), start=1):
            text = page.get_text()
            data = text.split('\n')
            data = list(filter(lambda x: x.strip(), data))
            all_data.extend(data)
            self.sinProgress.emit(self.filename, GuesserType.Text, idx, pages_count)
            time.sleep(0.01)
        c = Counter(all_data)
        items = c.most_common(self.most_common)
        # return items
        self.sinResult.emit(self.filename, items)
        doc.close()


class ImageGuesserThread(GuesserThread):
    sinNameRefMap = pyqtSignal(dict)

    name_ref_map = {}

    def run(self):
        self.name_ref_map = {}
        doc = fitz.open(self.filename)
        all_data = []
        pages_count = doc.page_count

        for idx, page in enumerate(doc.pages(), start=1):
            page.cleanContents()  # cleanup page painting commands
            xref = page.getContents()[0]  # get xref of the resulting source
            cont0 = doc.xrefStream(xref).decode()  # .splitlines()
            image_ids = re.findall(r'/Ima?g?e?(\d+) Do', cont0)
            for ele in image_ids:
                im = ImageGuesser.Image(ele)
                self.name_ref_map[str(im)] = ele
                all_data.append(str(im))
            time.sleep(0.01)
            self.sinProgress.emit(self.filename, GuesserType.Image, idx, pages_count)

        self.sinNameRefMap.emit(self.name_ref_map)
        c = Counter(all_data)
        items = c.most_common(self.most_common)
        self.sinResult.emit(self.filename, items)
        doc.close()


class WorkerThread(QThread):
    sinProgress = pyqtSignal(str, int, int, int)  # filename, idx in table, current, total
    sinDone = pyqtSignal(str, str, int)  # filename, output, idx in table
    # reset progress to show to doc progress
    sinReset = pyqtSignal(str, int, int, int)  # filename, idx in table, current, total
    # sin error
    sinError = pyqtSignal(str)

    def __init__(self, parent=None):
        super(WorkerThread, self).__init__(parent)
        self.filename = None
        self.index = None
        self.output = None
        self.text_list = []
        self.image_list = []
        self.to_doc = False

    def setVaribles(self, filename, idx, output, text_list, image_list, to_doc=False):
        self.filename = filename
        self.index = idx
        self.output = output
        self.text_list = text_list[:]
        self.image_list = image_list[:]
        self.to_doc = to_doc

    def run(self):
        raise NotImplementedError


class Algorithm(WorkerThread):
    desc = ''

    def remove_background(self, doc, page, sanitize=True):
        if sanitize:
            page.cleanContents(sanitize=sanitize)  # cleanup page painting commands
        xref = page.getContents()[0]  # get xref of the resulting source
        cont0 = doc.xrefStream(xref).decode().splitlines()  # and read it as lines of strings
        cont1 = []  # will contain reduced cont lines
        found = False  # indicates we are inside watermark instructions
        for line in cont0:
            if line.startswith("/Artifact") and "/Watermark" in line:  # start of watermark
                found = True  # switch on
                continue  # and skip line
            if found and line == "EMC":  # end of watermark
                found = False  # switch off
                continue  # and skip line
            if found is False:  # copy commands while outside watermarks
                cont1.append(line)
        cont = "\n".join(cont1)  # new paint commands source
        doc.updateStream(xref, cont.encode())  # replace old one with 'bytes' version

    def remove_text(self, page, text_list):
        if not text_list:
            return
        for text in text_list:
            areas = page.search_for(text)
            [page.addRedactAnnot(area, fill=(255, 255, 255)) for area in areas]
        page.apply_redactions()

    def remove_image(self, doc, image_list):
        for image in image_list:
            doc.updateStream(image, "".encode())

    def remove_image3(self, doc, page, image_list, sanitize=True):
        if not image_list:
            return
        pattern = r'[\d\. ]+cm[^(/Im)]+/Im({}) Do[^Q]+'.format('|'.join(image_list))
        pattern = re.compile(pattern)
        if sanitize:
            page.cleanContents(sanitize=sanitize)  # cleanup page painting commands
        xref = page.getContents()[0]  # get xref of the resulting source
        cont0 = doc.xrefStream(xref).decode()  # and read it as lines of strings
        cont = pattern.sub('', cont0)
        doc.updateStream(xref, cont.encode())

    @staticmethod
    def test_cleanContents(filename):
        flag = True
        try:
            doc = fitz.open(filename)
            page = doc.load_page(0)
            page.cleanContents()
            if fitz.TOOLS.mupdf_warnings(True):
                flag = False
            doc.close()
        except:
            pass

        return flag

    def remove_image2(self, doc, page, image_list):
        pattern = re.compile(r'/Ima?g?e?\d+ Do')
        pattern2 = re.compile('[\d\. ]+ cm')
        pattern3 = re.compile(r'/Im?g?e?(\d+) Do')

        page.cleanContents()  # cleanup page painting commands
        xref = page.getContents()[0]  # get xref of the resulting source
        cont0 = doc.xrefStream(xref).decode().splitlines()  # and read it as lines of strings

        matched = []
        for num, line in enumerate(cont0):
            # print(line)
            if pattern.match(line) and pattern3.findall(line)[0] in image_list:
                matched.append(num)

        excludes = []
        for idx in matched:
            start = idx
            while start > 0 and not pattern2.match(cont0[start]):
                start -= 1

            end = idx
            while cont0[end] != 'Q':
                end += 1

            if abs(end - start) > 10:
                print('Error')
                continue

            excludes.extend(range(start, end))

        includes = sorted(set(range(len(cont0))).difference(excludes))

        cont1 = []
        for i in includes:
            cont1.append(cont0[i])

        cont = "\n".join(cont1)

        doc.updateStream(xref, cont.encode())

    def __call__(self, filename, text_list, output, image_list):
        doc = fitz.open(filename)
        for page in doc.pages():
            self.remove_background(doc, page)
            page = doc.reload_page(page)
            self.remove_image3(doc, page, image_list)
            page = doc.reload_page(page)
            self.remove_text(page, text_list)
        # self.remove_image(doc, image_list)
        doc.save(output)

    def run(self):
        try:
            doc = fitz.open(self.filename)
            total = doc.page_count
            sanitize = self.test_cleanContents(self.filename)
            for num in range(total):
                idx = num + 1
                page = doc.load_page(num)
                # for idx, page in enumerate(doc.pages(), start=1):
                self.remove_background(doc, page, sanitize=sanitize)
                page = doc.reload_page(page)
                self.remove_image3(doc, page, self.image_list, sanitize=sanitize)
                page = doc.reload_page(page)
                self.remove_text(page, self.text_list)
                self.sinProgress.emit(self.filename, self.index, idx, total)
                time.sleep(0.01)
            # self.remove_image(doc, image_list)
            doc.save(self.output)
            if self.to_doc:
                self.sinReset.emit(self.filename, self.index, 0, total)
                self._to_doc(total)
        except:
            print(traceback.format_exc())

        self.sinDone.emit(self.filename, self.output, self.index)

    def _get_doc_filename(self, filename):
        filename = filename.replace('.pdf', '.docx') if filename.endswith('.pdf') else f'{filename}.docx'
        return filename

    def _to_doc(self, total):
        try:
            if not os.path.exists(self.output):
                return
            filename = self.output
            fp = open(filename, 'rb')
            parser = PDFParser(fp)
            doc = PDFDocument(parser)
            if not doc.is_extractable:
                raise PDFTextExtractionNotAllowed
            else:
                output = self._get_doc_filename(filename)
                rsrcmgr = PDFResourceManager()
                laparams = LAParams()
                device = PDFPageAggregator(rsrcmgr, laparams=laparams)
                interpreter = PDFPageInterpreter(rsrcmgr, device)

                document = Document()

                for idx, page in enumerate(PDFPage.create_pages(doc), start=1):
                    interpreter.process_page(page)
                    layout = device.get_result()
                    need_break = False
                    for x in layout:
                        if isinstance(x, LTTextBox):
                            need_break = True
                            document.add_paragraph(x.get_text().strip())
                    if need_break:
                        document.add_page_break()
                    self.sinProgress.emit(self.filename, self.index, idx, total)
                    time.sleep(0.01)
                document.save(output)
            fp.close()
        except:
            self.sinError.emit(traceback.format_exc())


class DefaultAlgorithm(Algorithm):
    desc = 'default'


file_info_dispatcher = {
    FileInfo.suffix: FileInfo,
    PDFInfo.suffix: PDFInfo
}

algorithm_dispatcher = {
    DefaultAlgorithm.desc: DefaultAlgorithm,
}


def get_file_info_cls(filename):
    ext = FileInfo(filename).filetype()
    if ext:
        return file_info_dispatcher.get(ext, FileInfo)
    return FileInfo


def get_algorithm_cls(algorithm):
    return algorithm_dispatcher.get(algorithm, DefaultAlgorithm)


class GuessDialog(QDialog):

    def __init__(self, parent=None):
        super(GuessDialog, self).__init__(parent)
        self.setWindowTitle('Guess Watermarks')
        self.file_label = QLabel(self)
        self.text_label = QLabel(self)
        self.image_label = QLabel(self)

        self.text_progress_bar = QProgressBar(self)
        self.image_progress_bar = QProgressBar(self)

        self.initUI()
        self.initSlots()

    def initUI(self):

        self.text_label.setText('Text Watermark')
        self.image_label.setText('Image Watermark')

        for bar in [self.text_progress_bar, self.image_progress_bar]:
            bar.setMinimum(0)
            bar.setOrientation(Qt.Horizontal)
            bar.setValue(0)

        layout = QGridLayout()
        layout.addWidget(self.file_label, 0, 0, 1, 3, Qt.AlignCenter)
        layout.addWidget(self.text_label, 1, 0, 1, 1)
        layout.addWidget(self.text_progress_bar, 1, 1, 1, 2)
        layout.addWidget(self.image_label, 2, 0, 1, 1)
        layout.addWidget(self.image_progress_bar, 2, 1, 1, 2)

        # fixed size for QGridLayout
        layout.setSizeConstraint(QLayout.SetFixedSize)
        self.setLayout(layout)

    def initSlots(self):
        pass

    def setFilename(self, filename):
        self.file_label.setText(filename)

    def setMax(self, value):
        for bar in [self.text_progress_bar, self.image_progress_bar]:
            bar.setMaximum(value)

    def updateImageProgress(self, value):
        self.image_progress_bar.setValue(value)

    def updateTextProgress(self, value):
        self.text_progress_bar.setValue(value)

    def reset(self, filename, max_value):
        self.file_label.setText(filename)
        for bar in [self.text_progress_bar, self.image_progress_bar]:
            bar.setMaximum(max_value)
            bar.setValue(0)

    def slotGuessProgress(self, gusser_type, value):
        if gusser_type == GuesserType.Text:
            self.updateTextProgress(value)
        elif gusser_type == GuesserType.Image:
            self.updateImageProgress(value)

        # check if done
        if self.text_progress_bar.value() == self.text_progress_bar.maximum() and \
                self.image_progress_bar.value() == self.image_progress_bar.maximum():
            self.accept()


class CallDialog(QDialog):
    FILENAME, PROGRESS, STATUS = range(3)

    sinDoneFile = pyqtSignal(str, str)  # input, output

    def __init__(self, parent=None):
        super(CallDialog, self).__init__(parent)
        self.setWindowTitle('In Progress')
        self.file_count = 0
        self.table = QTableView(self)
        self.model = self.create_model()
        self.total_progress = QProgressBar(self)

        self.initUI()
        self.initSlots()

    def initUI(self):
        self.table.setModel(self.model)

        # self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.NoSelection)

        self.table.horizontalHeader().setSectionResizeMode(self.FILENAME, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(self.PROGRESS, QHeaderView.Fixed)
        self.table.horizontalHeader().setSectionResizeMode(self.STATUS, QHeaderView.Fixed)
        # self.table.setColumnWidth(self.FILENAME, 300)
        self.table.setColumnWidth(self.PROGRESS, 100)
        self.table.setColumnWidth(self.STATUS, 50)
        self.table.setAlternatingRowColors(True)

        self.total_progress.setMinimum(0)
        self.total_progress.setValue(0)
        self.total_progress.setOrientation(Qt.Horizontal)
        self.total_progress.setMinimumWidth(500)

        layout = QVBoxLayout()
        layout.addWidget(self.table)
        layout.addWidget(self.total_progress)
        layout.setSizeConstraint(QLayout.SetFixedSize)

        self.setLayout(layout)

    def initSlots(self):
        pass

    def create_model(self):
        model = QStandardItemModel(0, 3, self)
        model.setHeaderData(self.FILENAME, Qt.Horizontal, 'Filename')
        model.setHeaderData(self.PROGRESS, Qt.Horizontal, 'Progress')
        model.setHeaderData(self.STATUS, Qt.Horizontal, 'Finish')
        return model

    def add_file(self, file):
        basename = os.path.basename(file)
        row = self.model.rowCount()
        self.model.insertRow(row)
        self.model.setData(self.model.index(row, self.FILENAME), basename)
        self.model.setData(self.model.index(row, self.PROGRESS), '')
        self.model.setData(self.model.index(row, self.STATUS), False)

    def reset_files(self, files):
        count = self.model.rowCount()
        self.model.removeRows(0, count)
        self.file_count = len(files)
        self.total_progress.setValue(0)
        self.total_progress.setMaximum(self.file_count)
        for file in files:
            self.add_file(file)

    def update_file(self, filename, idx, current, total):
        self.model.setData(self.model.index(idx, self.PROGRESS), f'{current} / {total}')

    def done_file(self, filename, output, idx):
        # emit to log view
        self.sinDoneFile.emit(filename, output)
        self.model.setData(self.model.index(idx, self.STATUS), True)
        value = self.total_progress.value()
        self.total_progress.setValue(value + 1)
        if self.total_progress.value() == self.total_progress.maximum():
            self.accept()


class App(QWidget):
    NAME, FORMAT, PAGES, PATH = range(4)
    WATERMARK, COUNT = range(2)

    sinGuessProgress = pyqtSignal(GuesserType, int)

    def __init__(self):
        super().__init__()
        self.title = 'Watermark Terminator - V2.1'
        self.pdf_doc = None  # pdf doc
        self.result_filename = None  # result doc filename
        self.left = 10
        self.top = 10
        self.width = 640
        self.height = 480
        # init ui components
        self.open_file_btn = QPushButton('File')
        self.open_folder_btn = QPushButton('Folder')
        self.algorithm_cobox = QComboBox()
        # self.save_chbox = QCheckBox('Save')
        self.preview_btn = QPushButton('Preview')
        self.run_selected_btn = QPushButton('Run Selected')
        self.run_all_btn = QPushButton('Run All')
        self.output_folder_ledit = QLineEdit()
        self.output_format_cobox = QComboBox()
        self.file_remove_btn = QPushButton('Remove')
        self.file_clear_btn = QPushButton('Clear')
        self.watermark_add_btn = QPushButton('Add')
        self.watermark_remove_btn = QPushButton('Remove')
        self.watermark_restore_btn = QPushButton('Restore')
        self.destination_btn = QPushButton('Destination:')
        self.file_list_view = QTreeView()
        self.watermark_view = QTreeView()

        self.cur_page = 0
        self.total_page = 0
        self.cur_page_label = QLabel('Current:')
        self.page_lineedit = QLineEdit()
        self.page_lineedit.setValidator(QIntValidator())
        self.total_page_label = QLabel('Total:')
        self.total_lineedit = QLineEdit()
        self.total_lineedit.setEnabled(False)
        self.prev_page_btn = QPushButton('Prev')
        self.next_page_btn = QPushButton('Next')

        self.source_view = QtImageViewer()
        self.target_view = QtImageViewer()
        self.log_view = QTextBrowser()
        self.clear_log_btn = QPushButton('Clear Log')
        # init models
        self.file_list_model = self.create_file_list_model(self)
        self.watermark_model = self.create_watermark_model(self)
        # global constant
        self.selected_file_row = -1
        self.guesser = WatermarkGuesser()
        # image guesser
        self.image_guesser = ImageGuesser()
        self.guessed_watermarks = []
        self.preview_filename = os.path.join(tempfile.gettempdir(), 'temporary_file_for_previewing.pdf')
        self.output_folder = ''
        self.output_filename_suffix = '_1.pdf'
        self.output_filename_doc_suffix = '_1.docx'
        # for thread
        self.name_ref_map = {}

        # progress dialogs
        self.guess_dialog = GuessDialog(self)
        # call dialog
        self.call_dialog = CallDialog(self)
        # init UI
        self.initUI()
        # init slots
        self.initSlots()

    def initUI(self):
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)

        # open
        open_group_box = QGroupBox('Open')
        open_group_layout = QGridLayout()
        open_group_layout.addWidget(self.open_file_btn, 0, 0)
        open_group_layout.addWidget(self.open_folder_btn, 1, 0)
        open_group_box.setLayout(open_group_layout)
        open_group_box.setSizePolicy(QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed))
        open_group_box.setMinimumHeight(100)

        # run
        self.algorithm_cobox.addItem('Default')
        run_group_box = QGroupBox('Run')
        run_group_layout = QGridLayout()
        run_group_layout.addWidget(QLabel('Algorithm:'), 0, 0)
        run_group_layout.addWidget(self.algorithm_cobox, 0, 1)
        run_group_layout.addWidget(self.preview_btn, 0, 2)
        run_group_layout.addWidget(self.run_selected_btn, 1, 0, 1, 2)
        run_group_layout.addWidget(self.run_all_btn, 1, 2)
        run_group_box.setLayout(run_group_layout)
        run_group_box.setSizePolicy(QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed))
        run_group_box.setMinimumHeight(100)

        # save
        self.output_format_cobox.addItem('PDF')
        self.output_format_cobox.addItem('DOCX')
        save_group_box = QGroupBox('Save')
        save_group_layout = QFormLayout()
        save_group_layout.addRow(QLabel('Format:'), self.output_format_cobox)
        save_group_layout.addRow(self.destination_btn, self.output_folder_ledit)
        save_group_box.setLayout(save_group_layout)
        self.output_folder_ledit.setSizePolicy(QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed))
        self.output_folder_ledit.setMinimumWidth(300)
        self.output_folder_ledit.setPlaceholderText('Default: Within Source Folder')
        self.output_folder_ledit.setEnabled(False)
        save_group_box.setSizePolicy(QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed))
        save_group_box.setMinimumHeight(100)

        file_list_box = QGroupBox('Files')
        file_list_layout = QGridLayout()
        # file list view
        self.file_list_view.setRootIsDecorated(False)
        self.file_list_view.setAlternatingRowColors(True)
        self.file_list_view.setModel(self.file_list_model)
        # disable editing
        self.file_list_view.setEditTriggers(QAbstractItemView.NoEditTriggers)
        file_list_layout.addWidget(self.file_list_view, 0, 0, 5, 2)
        file_list_layout.addWidget(self.file_remove_btn, 5, 0, 1, 1)
        file_list_layout.addWidget(self.file_clear_btn, 5, 1, 1, 1)

        file_list_box.setLayout(file_list_layout)

        # watermark view
        watermark_box = QGroupBox('Watermark')
        watermark_layout = QGridLayout()
        self.watermark_view.setRootIsDecorated(False)
        self.watermark_view.setAlternatingRowColors(True)
        self.watermark_view.setModel(self.watermark_model)
        watermark_layout.addWidget(self.watermark_view, 0, 0, 5, 3)

        watermark_layout.addWidget(self.watermark_add_btn, 5, 0, 1, 1)
        watermark_layout.addWidget(self.watermark_remove_btn, 5, 1, 1, 1)
        watermark_layout.addWidget(self.watermark_restore_btn, 5, 2, 1, 1)
        watermark_box.setLayout(watermark_layout)

        # pdf file view
        preview_box = QGroupBox('Preview')

        toolbar_layout = QGridLayout()
        toolbar_layout.addWidget(self.prev_page_btn, 0, 0)
        toolbar_layout.addWidget(self.cur_page_label, 0, 1)
        toolbar_layout.addWidget(self.page_lineedit, 0, 2)
        toolbar_layout.addWidget(self.total_page_label, 0, 3)
        toolbar_layout.addWidget(self.total_lineedit, 0, 4)
        toolbar_layout.addWidget(self.next_page_btn, 0, 5)

        pdf_view_layout = QHBoxLayout()
        pdf_view_layout.addWidget(self.source_view)
        pdf_view_layout.addWidget(self.target_view)

        preview_layout = QVBoxLayout()
        preview_layout.addLayout(toolbar_layout)
        preview_layout.addLayout(pdf_view_layout)

        preview_box.setLayout(preview_layout)

        top_layout = QHBoxLayout()
        top_layout.addWidget(open_group_box)
        top_layout.addWidget(run_group_box)
        top_layout.addWidget(save_group_box)
        top_layout.addStretch()

        middle_layout = QHBoxLayout()
        splitter = QSplitter()
        splitter.setChildrenCollapsible(False)
        splitter.addWidget(file_list_box)
        splitter.addWidget(watermark_box)
        splitter.addWidget(preview_box)
        middle_layout.addWidget(splitter)

        bottom_layout = QVBoxLayout()
        bottom_layout.addWidget(self.log_view)
        bottom_layout.addWidget(self.clear_log_btn)
        self.log_view.setSizePolicy(QSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed))
        self.log_view.setFixedHeight(100)

        # main layout
        main_layout = QVBoxLayout()
        main_layout.addLayout(top_layout)
        main_layout.addLayout(middle_layout)
        main_layout.addLayout(bottom_layout)
        self.setLayout(main_layout)

        self.show()

    def create_file_list_model(self, parent):
        model = QStandardItemModel(0, 4, parent)
        model.setHeaderData(self.NAME, Qt.Horizontal, 'Name')
        model.setHeaderData(self.FORMAT, Qt.Horizontal, 'Format')
        model.setHeaderData(self.PAGES, Qt.Horizontal, 'Pages')
        model.setHeaderData(self.PATH, Qt.Horizontal, 'PATH')
        return model

    def add_file(self, name, format, pages, path):
        self.file_list_model.insertRow(0)
        self.file_list_model.setData(self.file_list_model.index(0, self.NAME), name)
        self.file_list_model.setData(self.file_list_model.index(0, self.FORMAT), format)
        self.file_list_model.setData(self.file_list_model.index(0, self.PAGES), pages)
        self.file_list_model.setData(self.file_list_model.index(0, self.PATH), path)

    def create_watermark_model(self, parent):
        model = QStandardItemModel(0, 2, parent)
        model.setHeaderData(self.WATERMARK, Qt.Horizontal, 'Watermark')
        model.setHeaderData(self.COUNT, Qt.Horizontal, 'Count')
        return model

    def add_watermark(self, watermark, count=0):
        row = self.watermark_model.rowCount()
        self.watermark_model.insertRow(row)
        self.watermark_model.setData(self.watermark_model.index(row, self.WATERMARK), watermark)
        self.watermark_model.setData(self.watermark_model.index(row, self.COUNT), count)

    def initSlots(self):
        self.open_file_btn.clicked.connect(self.open_file)
        self.open_folder_btn.clicked.connect(self.open_folder)
        self.file_list_view.clicked.connect(self.change_selected_file)
        self.file_remove_btn.clicked.connect(self.remove_selected_files)
        self.file_clear_btn.clicked.connect(self.clear_files)
        self.watermark_add_btn.clicked.connect(self.new_watermark)
        self.watermark_remove_btn.clicked.connect(self.remove_watermark)
        self.watermark_restore_btn.clicked.connect(self.restore_watermark)
        self.preview_btn.clicked.connect(self.preview_result)
        self.run_selected_btn.clicked.connect(self.run_selected)
        self.run_all_btn.clicked.connect(self.run_all)
        self.destination_btn.clicked.connect(self.set_destination)
        self.next_page_btn.clicked.connect(self.next_page)
        self.prev_page_btn.clicked.connect(self.prev_page)
        self.page_lineedit.returnPressed.connect(self.set_page_number)
        self.sinGuessProgress.connect(self.guess_dialog.slotGuessProgress)
        # call progress
        self.call_dialog.sinDoneFile.connect(self.slotDoneFile)
        # clear log
        self.clear_log_btn.clicked.connect(self.clear_log_slot)

    def update_source_and_target_view(self):
        # update image
        self.load_pdf_page(self.source_view, self.cur_page)
        self.load_result_page(self.cur_page)

    @pyqtSlot()
    def set_page_number(self):
        if self.pdf_doc is None:
            return

        number = self.cur_page + 1
        try:
            number = int(self.page_lineedit.text())
        except:
            pass

        if number == self.cur_page + 1:
            return

        if number < 1 or number > self.total_page - 1:
            # update
            self.update_toolbar_display()
            return

        self.cur_page = number - 1
        self.update_toolbar_display()
        self.update_source_and_target_view()

    @pyqtSlot()
    def next_page(self):
        if not self.pdf_doc:
            return
        cur = self.cur_page
        total = self.total_page
        if cur + 1 >= total:
            return

        self.cur_page += 1
        self.update_toolbar_display()
        self.update_source_and_target_view()

    @pyqtSlot()
    def prev_page(self):
        if not self.pdf_doc:
            return
        cur = self.cur_page
        if cur - 1 < 0:
            return

        self.cur_page -= 1
        self.update_toolbar_display()
        self.update_source_and_target_view()

    def parse_file_attrs(self, filename):
        cls = get_file_info_cls(filename)
        ins = cls(filename)
        basename, format, pages, path = ins.basename(), ins.filetype(), ins.pages(), ins.fullpath()
        self.add_file(basename, format, pages, path)

    @pyqtSlot()
    def open_file(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        files, _ = QFileDialog.getOpenFileNames(self, "Open Files", "",
                                                "All Files (*);;PDF Files (*.pdf)", options=options)
        if files:
            try:
                [self.parse_file_attrs(file) for file in files]
            except:
                print(traceback.format_exc())
                self.log_exception(traceback.format_exc())

    @pyqtSlot()
    def open_folder(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        options |= QFileDialog.ShowDirsOnly
        folder = QFileDialog.getExistingDirectory(self, "Open Folder", "", options=options)
        if folder:
            try:
                for root, dirs, names in os.walk(folder):
                    [self.parse_file_attrs(os.path.join(root, name)) for name in names]
            except:
                print(traceback.format_exc())
                self.log_exception(traceback.format_exc())

    @pyqtSlot()
    def set_destination(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        folder = QFileDialog.getExistingDirectory(self, "Output Folder", "", options=options)
        if folder:
            self.output_folder_ledit.setText(folder)

    @staticmethod
    def clear_model(model):
        model.removeRows(0, model.rowCount())

    @staticmethod
    def remove_selected(view, model):
        indexes = view.selectedIndexes()
        rows = sorted(set(map(lambda x: x.row(), indexes)), reverse=True)
        [model.removeRow(row) for row in rows]
        # return removed rows
        return rows

    @pyqtSlot()
    def remove_selected_files(self):
        removed = self.remove_selected(self.file_list_view, self.file_list_model)
        if self.selected_file_row in removed:
            # reset pdf view
            self.reset_pdf_view(True, True)
            self.selected_file_row = -1
            self.guessed_watermarks = []
            self.reset_watermarks()

    @pyqtSlot()
    def clear_files(self):
        self.clear_model(self.file_list_model)
        # reset watermarks
        self.reset_watermarks()
        # reset pdf view
        self.reset_pdf_view(True, True)
        # reset log view
        self.clear_log_slot()

    def reset_pdf_view(self, source=True, target=False):
        if source:
            self.source_view.clearImage()
        if target:
            self.target_view.clearImage()

    def reset_watermarks(self):
        self.clear_model(self.watermark_model)
        self.guessed_watermarks = []

    @pyqtSlot()
    def new_watermark(self):
        count = self.watermark_model.rowCount()
        self.watermark_model.insertRow(count)
        self.watermark_model.setData(self.watermark_model.index(count, self.WATERMARK), 'Tap to edit')
        self.watermark_model.setData(self.watermark_model.index(count, self.COUNT), 'unknown')

    @pyqtSlot()
    def remove_watermark(self):
        self.remove_selected(self.watermark_view, self.watermark_model)

    @pyqtSlot()
    def restore_watermark(self):
        self.clear_model(self.watermark_model)
        watermarks = self.guessed_watermarks[:]
        [self.add_watermark(*ele) for ele in watermarks]

    def update_watermark(self):
        if self.pdf_doc is None:
            return
        self.guesser.doc = self.pdf_doc
        watermarks = self.guesser()
        # image watermarks
        self.image_guesser.doc = self.pdf_doc
        image_watermarks = self.image_guesser()
        self.guessed_watermarks = watermarks + image_watermarks
        [self.add_watermark(*ele) for ele in self.guessed_watermarks[::-1]]

    def load_as_image(self, page):
        pixmap = page.get_pixmap(mat=fitz.Matrix(2, 2))
        data = pixmap.getPNGData()
        image = QImage()
        image.loadFromData(data)
        return image

    def load_pdf_page(self, viewer, page_number=0, filename=None):
        doc = self.pdf_doc
        if filename is not None:
            doc = fitz.open(filename)

        if doc is None:
            return
        page = doc.load_page(page_number)
        image = self.load_as_image(page)
        viewer.setImage(image)

    def load_result_page(self, page_number=0, filename=None):
        if filename is None:
            filename = self.result_filename
        else:
            # update result doc filename
            self.result_filename = filename

        if filename is not None and os.path.exists(filename):
            doc = fitz.open(filename)

            page = doc.load_page(page_number)
            image = self.load_as_image(page)
            self.target_view.setImage(image)

            doc.close()

    @pyqtSlot()
    def change_selected_file(self):
        # get selected file
        index = self.file_list_view.selectedIndexes()
        # get filename
        filename = index[self.PATH].data()
        # update constant
        self.selected_file_row = index[self.PATH].row()
        # is pdf
        is_pdf = PDFInfo.is_pdf(filename)
        # update preview
        self.reset_pdf_view(True, True)
        self.reset_watermarks()
        if is_pdf:
            # start image watermark guesser
            image_guesser_thread = ImageGuesserThread(self)
            image_guesser_thread.filename = filename
            image_guesser_thread.sinProgress.connect(self.slotProgress)
            image_guesser_thread.sinResult.connect(self.slotResult)
            image_guesser_thread.sinNameRefMap.connect(self.slotNameRefMap)
            image_guesser_thread.finished.connect(image_guesser_thread.deleteLater)

            image_guesser_thread.start()

            # start text watermark guesser
            text_guesser_thread = TextGuesserThread(self)
            text_guesser_thread.filename = filename
            text_guesser_thread.sinProgress.connect(self.slotProgress)
            text_guesser_thread.sinResult.connect(self.slotResult)
            text_guesser_thread.finished.connect(text_guesser_thread.deleteLater)
            text_guesser_thread.start()

            self.pdf_doc = fitz.open(filename)
            self.result_filename = None
            self.total_page = self.pdf_doc.page_count
            self.cur_page = 0
            self.load_pdf_page(self.source_view)
            # self.update_watermark()
            self.update_toolbar_display()
            # guess dialog
            self.guess_dialog.reset(filename, self.total_page)
            self.guess_dialog.exec()

    def update_toolbar_display(self):
        self.page_lineedit.setText(f'{self.cur_page + 1}')
        self.total_lineedit.setText(f'{self.total_page}')

    def get_watermark_list(self):
        count = self.watermark_model.rowCount()
        data = [self.watermark_model.data(self.watermark_model.index(i, self.WATERMARK)) for i in range(count)]
        image_map = self.name_ref_map
        text = [i for i in data if i not in image_map]
        image = [image_map.get(i) for i in data if i in image_map]
        return text, image

    def get_selected_file(self):
        index = self.file_list_view.selectedIndexes()
        if index:
            filename = index[self.PATH].data()
            return filename

    def need_to_doc(self):
        return self.output_format_cobox.currentIndex() == 1

    def single_run(self, _input, output, update_preview=False):
        text, image = self.get_watermark_list()

        self.call_dialog.reset_files([_input])
        worker_thread = Algorithm(self)
        worker_thread.setVaribles(_input, 0, output, text, image, self.need_to_doc())
        worker_thread.sinProgress.connect(self.call_dialog.update_file)
        worker_thread.sinReset.connect(self.call_dialog.update_file)
        worker_thread.sinDone.connect(self.call_dialog.done_file)
        worker_thread.sinError.connect(self.log_exception)
        worker_thread.finished.connect(worker_thread.deleteLater)
        worker_thread.start()

        if self.call_dialog.exec() == QDialog.Accepted:
            if update_preview:
                # load preview
                self.load_result_page(self.cur_page, filename=output)

    @pyqtSlot()
    def preview_result(self):
        _input = self.get_selected_file()
        output = self.preview_filename
        # remove if exists
        try:
            if os.path.exists(output):
                os.remove(output)
        except:
            self.log_exception(traceback.format_exc())
        self.single_run(_input, output, update_preview=True)

    def _make_output(self, _input, template):
        if not _input:
            return _input
        basename = os.path.basename(_input)
        name, ext = os.path.splitext(basename)
        output_folder = self.output_folder_ledit.text() or os.path.dirname(_input)
        return os.path.join(output_folder, f'{name}{template}')

    def gen_output(self, _input):
        return self._make_output(_input, self.output_filename_suffix)

    def get_doc_output(self, _input):
        return self._make_output(_input, self.output_filename_doc_suffix)

    @pyqtSlot()
    def run_selected(self):
        _input = self.get_selected_file()
        output = self.gen_output(_input)
        self.single_run(_input, output, update_preview=True)

    def get_all_files(self):
        count = self.file_list_model.rowCount()
        paths = [self.file_list_model.data(self.file_list_model.index(i, self.PATH)) for i in range(count)]
        return paths

    @pyqtSlot()
    def run_all(self):
        files = self.get_all_files()
        self.call_dialog.reset_files(files)
        text, image = self.get_watermark_list()
        for idx, file in enumerate(files):
            output = self.gen_output(file)
            worker_thread = Algorithm(self)
            worker_thread.setVaribles(file, idx, output, text, image, self.need_to_doc())
            worker_thread.sinProgress.connect(self.call_dialog.update_file)
            worker_thread.sinReset.connect(self.call_dialog.update_file)
            worker_thread.sinDone.connect(self.call_dialog.done_file)
            worker_thread.sinError.connect(self.log_exception)
            worker_thread.finished.connect(worker_thread.deleteLater)
            worker_thread.start()

        if self.call_dialog.exec() == QDialog.Accepted:
            # update target view using processed file
            _input = self.get_selected_file()
            output = self.gen_output(_input)
            if output:
                self.load_result_page(self.cur_page, filename=output)

    def slotDoneFile(self, input, output):
        doc_output = self.get_doc_output(input) if self.need_to_doc() else None
        self.log_process(input, output, doc_output=doc_output)

    def log_process(self, input, output, doc_output=None):
        if output == self.preview_filename:
            return
        data = f'Input :\t{input}\n' \
               f'Output:\t{output}\n'
        if doc_output and os.path.exists(doc_output):
            data = f'{data}' \
                   f'Docx Output: {doc_output}\n'
        self.log_view.append(data)

    def log_exception(self, exception):
        self.log_view.append('=' * 10 + 'Exception' + '=' * 10 + '\n')
        self.log_view.append(exception + '\n')
        self.log_view.append('=' * (10 * 2 + len('Exception')) + '\n')

    def get_output_format(self):
        return self.output_format_cobox.currentText()

    def slotProgress(self, filename, guesser_type, current, total):
        self.sinGuessProgress.emit(guesser_type, current)

    def slotResult(self, filename, items):
        [self.add_watermark(*ele) for ele in items]
        self.guessed_watermarks.extend(items)

    def slotNameRefMap(self, d):
        self.name_ref_map = d

    @pyqtSlot()
    def clear_log_slot(self):
        self.log_view.clear()


if __name__ == '__main__':
    try:
        import popplerqt5
    except:
        import traceback

        print(traceback.format_exc())
    app = QApplication(sys.argv)
    ex = App()
    sys.exit(app.exec_())
